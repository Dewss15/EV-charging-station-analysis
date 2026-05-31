"""
CE634 Data Mining and Data Warehousing — Report Generator
Generates a professional .docx report for the EV Charging Station Analysis case study.
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── paths ──────────────────────────────────────────────────────────────────────
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
IMG = {
    "correlation":    os.path.join(BASE_DIR, "correlation.png"),
    "euclidean":      os.path.join(BASE_DIR, "euclidean.png"),
    "confusion":      os.path.join(BASE_DIR, "confusion_matrix.png"),
    "feature_imp":    os.path.join(BASE_DIR, "feature_imp.png"),
    "class_dist":     os.path.join(BASE_DIR, "class.png"),
    "dashboard":      os.path.join(BASE_DIR, "final_dashboard.png"),
}
OUTPUT = os.path.join(BASE_DIR, "CE634_EV_Charging_Case_Study_Report.docx")

# ── colour palette ─────────────────────────────────────────────────────────────
DARK_BLUE   = RGBColor(0x1A, 0x37, 0x5E)   # headings
MID_BLUE    = RGBColor(0x27, 0x5C, 0x9E)   # sub-headings
ACCENT      = RGBColor(0x00, 0x78, 0xD4)   # table header rows
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GRAY  = RGBColor(0xF2, 0xF5, 0xF9)
BODY_COLOR  = RGBColor(0x1E, 0x1E, 0x1E)
CAPTION_CLR = RGBColor(0x55, 0x55, 0x55)

# ── helpers ────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, rgb_hex):
    """Set table cell background colour."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), rgb_hex)
    tcPr.append(shd)


def add_page_number(doc):
    """Insert page number field in the footer."""
    section = doc.sections[0]
    footer  = section.footer
    para    = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def add_horizontal_rule(doc):
    """Add a thin horizontal rule paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1A375E")
    pBdr.append(bottom)
    pPr.append(pBdr)


def body_para(doc, text, bold=False, italic=False, indent=False):
    """Add a well-formatted body paragraph."""
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if indent:
        p.paragraph_format.left_indent = Cm(0.8)
    run = p.add_run(text)
    run.font.size   = Pt(11)
    run.font.color.rgb = BODY_COLOR
    run.bold        = bold
    run.italic      = italic
    return p


def bullet(doc, text, level=0):
    """Add a bullet-list paragraph."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Cm(0.6 + level * 0.5)
    run = p.add_run(text)
    run.font.size  = Pt(11)
    run.font.color.rgb = BODY_COLOR
    return p


def figure_caption(doc, number, title):
    """Add a figure caption below an image."""
    p   = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(12)
    run = p.add_run(f"Figure {number}: {title}")
    run.font.size   = Pt(9.5)
    run.font.italic = True
    run.font.color.rgb = CAPTION_CLR


def table_caption(doc, number, title):
    """Add a table caption above a table."""
    p   = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(f"Table {number}: {title}")
    run.font.size   = Pt(9.5)
    run.font.bold   = True
    run.font.color.rgb = CAPTION_CLR


def insert_image(doc, path, width=Inches(5.5), caption_num=None, caption_title=""):
    """Insert an image centred with an optional caption."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    if os.path.exists(path):
        run = p.add_run()
        run.add_picture(path, width=width)
    else:
        run = p.add_run(f"[Image not found: {os.path.basename(path)}]")
        run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    if caption_num:
        figure_caption(doc, caption_num, caption_title)


def heading1(doc, text):
    h = doc.add_heading(text, level=1)
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after  = Pt(6)
    for run in h.runs:
        run.font.color.rgb = DARK_BLUE
        run.font.size      = Pt(14)
        run.font.bold      = True
    return h


def heading2(doc, text):
    h = doc.add_heading(text, level=2)
    h.paragraph_format.space_before = Pt(10)
    h.paragraph_format.space_after  = Pt(4)
    for run in h.runs:
        run.font.color.rgb = MID_BLUE
        run.font.size      = Pt(12)
        run.font.bold      = True
    return h


def styled_table(doc, headers, rows, col_widths=None):
    """Create a nicely styled table with a blue header row."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_bg(hdr_cells[i], "275C9E")
        for para in hdr_cells[i].paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.bold       = True
                run.font.color.rgb  = WHITE
                run.font.size       = Pt(10)

    # data rows
    for r_idx, row_data in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        bg = "F2F5F9" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            row_cells[c_idx].text = val
            row_cells[c_idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_bg(row_cells[c_idx], bg)
            for para in row_cells[c_idx].paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.size       = Pt(10)
                    run.font.color.rgb  = BODY_COLOR

    # column widths
    if col_widths:
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                cell.width = col_widths[i]

    return table


# ══════════════════════════════════════════════════════════════════════════════
# MAIN BUILDER
# ══════════════════════════════════════════════════════════════════════════════

def build_report():
    doc = Document()

    # ── page setup ─────────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width  = Inches(8.27)   # A4
    section.page_height = Inches(11.69)
    section.left_margin   = Inches(1.0)
    section.right_margin  = Inches(1.0)
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    add_page_number(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # 1. TITLE PAGE
    # ══════════════════════════════════════════════════════════════════════════
    for _ in range(4):
        doc.add_paragraph()

    # Institution
    inst = doc.add_paragraph()
    inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = inst.add_run("Padre Conceicao College of Engineering")
    r.font.size = Pt(13); r.font.bold = True; r.font.color.rgb = DARK_BLUE

    doc.add_paragraph()

    course_p = doc.add_paragraph()
    course_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = course_p.add_run("CE634: Data Mining and Data Warehousing")
    r.font.size = Pt(12); r.font.color.rgb = MID_BLUE; r.font.bold = True

    for _ in range(2):
        doc.add_paragraph()

    # Main title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title_p.add_run("Predictive Analysis of EV Charging Demand,\nGrid Load Dynamics, and Peak Load Risk")
    r.font.size = Pt(22); r.font.bold = True; r.font.color.rgb = DARK_BLUE

    doc.add_paragraph()
    add_horizontal_rule(doc)
    doc.add_paragraph()

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub_p.add_run("Case Study Report")
    r.font.size = Pt(13); r.font.color.rgb = MID_BLUE; r.font.italic = True

    for _ in range(3):
        doc.add_paragraph()

    # Students
    students = [
        ("Dewpearl Gonsalves", "23CE114"),
        ("Shail Joshi",        "23CE163"),
        ("Flyson Dias",        "23CE118"),
    ]
    for name, roll in students:
        sp = doc.add_paragraph()
        sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = sp.add_run(f"{name}  •  Roll No. {roll}")
        r.font.size = Pt(11.5); r.font.color.rgb = BODY_COLOR

    for _ in range(2):
        doc.add_paragraph()

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = date_p.add_run("Academic Year: 2025–2026  |  May 2026")
    r.font.size = Pt(10.5); r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 2. TABLE OF CONTENTS (static)
    # ══════════════════════════════════════════════════════════════════════════
    toc_heading = doc.add_paragraph()
    toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = toc_heading.add_run("Table of Contents")
    r.font.size = Pt(16); r.font.bold = True; r.font.color.rgb = DARK_BLUE
    toc_heading.paragraph_format.space_after = Pt(12)

    add_horizontal_rule(doc)
    doc.add_paragraph()

    toc_entries = [
        ("Abstract",                                        "3"),
        ("1. Introduction",                                 "3"),
        ("2. Problem Statement",                            "4"),
        ("3. Dataset Description",                          "4"),
        ("4. Tools and Technologies",                       "4"),
        ("5. Data Understanding and Preprocessing",         "5"),
        ("6. Similarity and Dissimilarity Analysis",        "6"),
        ("   6.1 Correlation Matrix",                       "6"),
        ("   6.2 Euclidean Distance Analysis",              "6"),
        ("7. Classification Analysis",                      "7"),
        ("8. Feature Importance Analysis",                  "8"),
        ("9. Class Distribution Analysis",                  "9"),
        ("10. Clustering Analysis",                         "9"),
        ("11. Outlier Detection",                          "10"),
        ("12. Dashboard Visualization",                    "10"),
        ("13. Results and Discussion",                     "11"),
        ("14. Conclusion",                                 "12"),
        ("References",                                     "13"),
    ]

    for entry, page in toc_entries:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(2)
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(5.5), leader="…" if False else None)
        run_left  = p.add_run(entry)
        run_left.font.size = Pt(10.5)
        if not entry.startswith("   "):
            run_left.font.bold = True
            run_left.font.color.rgb = DARK_BLUE
        else:
            run_left.font.color.rgb = MID_BLUE
        run_dots = p.add_run(" " + "." * max(1, 70 - len(entry) * 2) + " " + page)
        run_dots.font.size = Pt(10.5)
        run_dots.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # ABSTRACT
    # ══════════════════════════════════════════════════════════════════════════
    heading1(doc, "Abstract")
    body_para(doc,
        "This case study presents a comprehensive data mining and warehousing analysis "
        "of an EV (Electric Vehicle) charging station dataset comprising 2,800 records "
        "and 10 attributes sourced from Kaggle. The study undertakes an end-to-end "
        "machine learning pipeline encompassing data preprocessing, feature engineering, "
        "similarity and dissimilarity analysis, supervised classification, unsupervised "
        "clustering, and anomaly detection."
    )
    body_para(doc,
        "A Random Forest Classifier was trained to predict the categorical target variable "
        "peak_load_risk (Low, Medium, High), achieving an overall accuracy of 45.18% after "
        "addressing class imbalance through SMOTE. Pearson correlation analysis revealed "
        "predominantly weak linear relationships among numerical features, while Euclidean "
        "distance computation highlighted notable dissimilarities across city zones. "
        "K-Means clustering (k = 3) identified three distinct operational charging profiles, "
        "and Isolation Forest detected 84 anomalous charging events representing irregular "
        "combinations of energy demand and grid load. The findings offer actionable insights "
        "for EV infrastructure planning and smart grid load management."
    )
    add_horizontal_rule(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # 1. INTRODUCTION
    # ══════════════════════════════════════════════════════════════════════════
    heading1(doc, "1. Introduction")
    body_para(doc,
        "The global transition towards electric mobility has accelerated the deployment of EV "
        "charging infrastructure across urban and suburban environments. As the penetration of "
        "electric vehicles continues to rise, the electric grid faces unprecedented challenges "
        "in managing fluctuating and often unpredictable charging loads. Unlike traditional "
        "household appliances, EV chargers—particularly fast-charging units—introduce high-power "
        "draw events that can destabilize local distribution networks if left unmanaged."
    )
    body_para(doc,
        "Charging station analytics, powered by data mining techniques, provide operators with "
        "the ability to understand historical usage patterns, anticipate demand surges, and "
        "allocate grid resources efficiently. By mining temporal, geographic, and operational "
        "features from station-level data, utilities can shift from reactive grid management "
        "to a proactive, intelligence-driven operational model."
    )
    body_para(doc,
        "Grid load management is a critical concern in regions with high EV adoption rates. "
        "Peak demand events—where simultaneous charging sessions exceed the safe operational "
        "threshold of a grid segment—can lead to voltage instability, equipment stress, and "
        "potential service interruptions. Accurate prediction of peak load risk levels enables "
        "grid operators to preemptively engage demand-response programmes, deploy battery "
        "buffers, or throttle charging rates at specific stations."
    )
    body_para(doc,
        "This case study applies data mining methodologies to the EV Charging Station Usage "
        "and Grid Load Analysis dataset to build a predictive model for peak load risk "
        "classification, identify behavioural charging clusters, and detect anomalous high-stress "
        "events that may compromise grid reliability."
    )
    add_horizontal_rule(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # 2. PROBLEM STATEMENT
    # ══════════════════════════════════════════════════════════════════════════
    heading1(doc, "2. Problem Statement")
    body_para(doc,
        "The primary objective of this study is to apply data mining techniques to analyse "
        "EV charging station behaviour and its impact on grid load. Specifically, the study "
        "aims to:"
    )
    bullet(doc, "Analyse EV charging behaviour across different city zones and station types.")
    bullet(doc, "Study temporal and spatial patterns in grid load.")
    bullet(doc, "Predict the level of peak load risk (Low, Medium, or High) using a supervised classification model.")
    bullet(doc, "Identify distinct operational charging clusters through unsupervised learning.")
    bullet(doc, "Detect anomalous charging events that may represent unusual or potentially harmful grid stress scenarios.")
    add_horizontal_rule(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # 3. DATASET DESCRIPTION
    # ══════════════════════════════════════════════════════════════════════════
    heading1(doc, "3. Dataset Description")
    body_para(doc,
        "The dataset used in this study is the EV Charging Station Usage and Grid Load Analysis "
        "dataset, obtained from Kaggle. It simulates real-world EV station operational data "
        "across multiple city zones and station types."
    )

    table_caption(doc, 1, "Dataset Summary Statistics")
    styled_table(doc,
        headers=["Property", "Value"],
        rows=[
            ["Total Records",    "2,800"],
            ["Total Attributes", "10"],
            ["Missing Values",   "None"],
            ["Duplicate Records","None"],
            ["Target Variable",  "peak_load_risk (Low / Medium / High)"],
        ],
        col_widths=[Inches(2.8), Inches(3.2)],
    )
    doc.add_paragraph()

    body_para(doc, "The dataset contains the following features:", bold=True)
    table_caption(doc, 2, "Dataset Feature Descriptions")
    styled_table(doc,
        headers=["Feature", "Type", "Description"],
        rows=[
            ["record_id",                       "Integer",     "Unique identifier for each charging session"],
            ["date_time",                        "DateTime",    "Timestamp of the charging session"],
            ["city_zone",                        "Categorical", "Geographic zone (North, South, East, West)"],
            ["station_type",                     "Categorical", "Type of charger (Fast, Slow, Ultra-fast)"],
            ["vehicles_charged",                 "Integer",     "Number of vehicles charged in a session"],
            ["avg_charging_duration_minutes",    "Float",       "Average session duration in minutes"],
            ["energy_dispensed_kwh",             "Float",       "Total energy delivered (kWh)"],
            ["grid_load_mw",                     "Float",       "Grid load drawn during session (MW)"],
            ["renewable_energy_used_percent",    "Float",       "Proportion of renewable energy used (%)"],
            ["peak_load_risk",                   "Categorical", "Risk classification: Low / Medium / High"],
        ],
        col_widths=[Inches(2.2), Inches(1.1), Inches(2.7)],
    )
    doc.add_paragraph()
    add_horizontal_rule(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # 4. TOOLS AND TECHNOLOGIES
    # ══════════════════════════════════════════════════════════════════════════
    heading1(doc, "4. Tools and Technologies")
    body_para(doc,
        "The following technologies and libraries were employed throughout the data mining "
        "pipeline:"
    )

    table_caption(doc, 3, "Tools and Libraries Used")
    styled_table(doc,
        headers=["Tool / Library", "Purpose"],
        rows=[
            ["Python 3.x",       "Primary programming language for all analysis and modelling"],
            ["Pandas",           "Data loading, manipulation, and exploratory data analysis"],
            ["NumPy",            "Numerical computations and array operations"],
            ["Scikit-Learn",     "Machine learning models: Random Forest, K-Means, Isolation Forest, SMOTE, StandardScaler"],
            ["Matplotlib",       "Plotting and static visualisation"],
            ["Seaborn",          "Statistical data visualisation (heatmaps, distribution plots)"],
            ["Google Colab",     "Cloud-based Jupyter notebook environment for reproducible experimentation"],
            ["VS Code",          "Local development environment for script editing and debugging"],
        ],
        col_widths=[Inches(1.8), Inches(4.2)],
    )
    doc.add_paragraph()
    add_horizontal_rule(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # 5. DATA UNDERSTANDING AND PREPROCESSING
    # ══════════════════════════════════════════════════════════════════════════
    heading1(doc, "5. Data Understanding and Preprocessing")

    heading2(doc, "5.1 Dataset Inspection")
    body_para(doc,
        "The raw dataset was loaded from the Kaggle CSV file. An initial inspection confirmed "
        "2,800 records with 10 attributes. No missing values were detected in any column, "
        "and no duplicate rows were found, indicating a clean source dataset suitable for "
        "direct analysis without imputation."
    )

    heading2(doc, "5.2 Feature Engineering")
    body_para(doc,
        "The date_time column was parsed into a Python datetime object, enabling the extraction "
        "of two temporal features:"
    )
    bullet(doc, "hour_of_day — The hour (0–23) extracted from the timestamp, capturing intraday charging patterns and peak-hour effects.")
    bullet(doc, "day_of_week — The integer day of the week (0 = Monday, 6 = Sunday), capturing weekly behavioral cycles.")
    body_para(doc,
        "These engineered features are critical for temporal pattern recognition, as EV "
        "charging demand is strongly correlated with daily human activity schedules."
    )

    heading2(doc, "5.3 Categorical Encoding")
    body_para(doc,
        "The categorical columns station_type and city_zone were encoded using One-Hot Encoding "
        "with drop_first=True to avoid multicollinearity (dummy variable trap). This produced "
        "binary indicator columns for each category level, allowing the model to interpret "
        "categorical identity without imposing an ordinal ranking."
    )

    heading2(doc, "5.4 Target Variable Mapping")
    body_para(doc,
        "The target variable peak_load_risk was mapped from its string representation to "
        "ordinal integer labels as follows: Low → 0, Medium → 1, High → 2. This ordinal "
        "encoding preserves the inherent risk severity ordering and is compatible with "
        "standard scikit-learn classifiers."
    )

    heading2(doc, "5.5 Feature Scaling")
    body_para(doc,
        "StandardScaler was applied to all numerical features (energy_dispensed_kwh, "
        "avg_charging_duration_minutes, grid_load_mw, vehicles_charged, "
        "renewable_energy_used_percent). Standardisation ensures that features with "
        "different units and magnitudes contribute equally to distance-based computations "
        "(Euclidean distance, K-Means) and improves convergence for ensemble models."
    )

    heading2(doc, "5.6 Final Dataset")
    body_para(doc,
        "Following all preprocessing steps, the final dataset dimensions were (2,800, 15), "
        "comprising the original numerical features, two temporal features, and the one-hot "
        "encoded categorical columns. The record_id identifier column was excluded from "
        "model training to prevent data leakage."
    )

    table_caption(doc, 4, "Preprocessing Steps Summary")
    styled_table(doc,
        headers=["Step", "Technique", "Reason"],
        rows=[
            ["1", "Null Check",              "Confirm data completeness"],
            ["2", "Duplicate Removal",       "Ensure sample independence"],
            ["3", "Datetime Parsing",        "Enable temporal feature extraction"],
            ["4", "Feature Engineering",     "Capture time-of-day and day-of-week patterns"],
            ["5", "One-Hot Encoding",        "Convert categorical variables for ML compatibility"],
            ["6", "Ordinal Target Mapping",  "Encode target labels for classification"],
            ["7", "StandardScaler",          "Normalise feature magnitudes"],
            ["8", "Drop record_id",          "Prevent data leakage"],
        ],
        col_widths=[Inches(0.4), Inches(2.0), Inches(3.6)],
    )
    doc.add_paragraph()
    add_horizontal_rule(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # 6. SIMILARITY AND DISSIMILARITY ANALYSIS
    # ══════════════════════════════════════════════════════════════════════════
    heading1(doc, "6. Similarity and Dissimilarity Analysis")

    heading2(doc, "6.1 Correlation Matrix")
    body_para(doc,
        "A Pearson correlation matrix was computed across all numerical features to quantify "
        "the strength of linear relationships. The heatmap (Figure 1) visualises these "
        "pairwise correlation coefficients."
    )
    body_para(doc,
        "The analysis revealed predominantly weak linear relationships among all numerical "
        "variables, with most Pearson correlation values falling close to zero. No single "
        "numerical feature exhibits a strong correlation with another. This near-absence of "
        "linear correlation has important implications: it suggests that the features are "
        "relatively independent of one another, which can be beneficial for ensemble models "
        "but may also indicate that peak load risk is governed by complex, non-linear "
        "interactions among features rather than by any single dominant predictor. This "
        "observation partly explains the moderate classification accuracy observed in "
        "subsequent modelling experiments."
    )
    insert_image(doc, IMG["correlation"], width=Inches(5.2),
                 caption_num=1, caption_title="Pearson Correlation Heatmap of Numerical Features")

    heading2(doc, "6.2 Euclidean Distance Analysis")
    body_para(doc,
        "Euclidean distances were computed between the mean feature vectors of each city zone "
        "to quantify operational dissimilarity across geographic areas. The resulting distance "
        "matrix is visualised as a heatmap in Figure 2."
    )
    body_para(doc,
        "Key findings from the Euclidean distance analysis:"
    )
    bullet(doc, "South and West zones are the most operationally similar (distance ≈ 0.85), indicating comparable charging load profiles, likely driven by similar population densities or commuting patterns.")
    bullet(doc, "East and South zones are the most dissimilar (distance ≈ 3.64), suggesting fundamentally different charging behaviours, possibly due to differences in station types, vehicle density, or grid infrastructure.")
    body_para(doc,
        "These dissimilarity findings reinforce the importance of zone-specific grid management "
        "strategies rather than a uniform nationwide approach. Grid operators can prioritise "
        "infrastructure upgrades in zones exhibiting the highest operational divergence."
    )
    insert_image(doc, IMG["euclidean"], width=Inches(4.8),
                 caption_num=2, caption_title="Euclidean Distance Heatmap Across City Zones")
    add_horizontal_rule(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # 7. CLASSIFICATION ANALYSIS
    # ══════════════════════════════════════════════════════════════════════════
    heading1(doc, "7. Classification Analysis")

    heading2(doc, "7.1 Methodology")
    body_para(doc,
        "A Random Forest Classifier was selected for peak load risk prediction due to its "
        "robustness against overfitting, ability to handle mixed feature types, and inherent "
        "support for feature importance quantification. The following pipeline was implemented:"
    )
    bullet(doc, "Train-Test Split: 80% training, 20% testing with stratified sampling to preserve class proportions.")
    bullet(doc, "SMOTE (Synthetic Minority Over-sampling Technique): Applied to the training set to address class imbalance by generating synthetic samples for the Medium and High risk minority classes.")
    bullet(doc, "Random Forest Training: 100 decision tree estimators with default hyperparameters.")
    bullet(doc, "Performance Evaluation: Accuracy, precision, recall, F1-score, and confusion matrix.")

    heading2(doc, "7.2 Results")
    body_para(doc, "Overall Accuracy: 45.18%", bold=True)
    doc.add_paragraph()

    table_caption(doc, 5, "Classification Report — Random Forest Classifier")
    styled_table(doc,
        headers=["Class", "Precision", "Recall", "F1-Score", "Support"],
        rows=[
            ["Low (0)",    "0.55", "0.62", "0.58", "281"],
            ["Medium (1)", "0.34", "0.33", "0.33", "166"],
            ["High (2)",   "0.30", "0.21", "0.25", "113"],
            ["Weighted Avg.", "0.43", "0.45", "0.44", "560"],
        ],
        col_widths=[Inches(1.6), Inches(0.9), Inches(0.9), Inches(0.9), Inches(0.9)],
    )
    doc.add_paragraph()

    heading2(doc, "7.3 Interpretation")
    body_para(doc,
        "The classification results reflect the inherent complexity of predicting peak load risk "
        "from the available features:"
    )
    bullet(doc, "Low-risk events (F1 = 0.58) are predicted most accurately, likely because this majority class provides more training signal and exhibits clearer feature boundaries.")
    bullet(doc, "Medium-risk events (F1 = 0.33) are the most difficult to classify, as they represent a transitional risk state with overlapping feature characteristics shared with both Low and High classes.")
    bullet(doc, "High-risk events (F1 = 0.25) suffer from recall limitations, indicating that the model frequently misclassifies High-risk events as Medium or Low, which is the most operationally critical error type.")
    bullet(doc, "SMOTE improved the model's ability to detect minority class instances compared to the baseline imbalanced model, but the fundamental complexity of non-linear risk boundaries limits achievable accuracy.")
    body_para(doc,
        "The moderate accuracy is consistent with the weak linear correlations observed in the "
        "similarity analysis and is expected given the synthetic nature of the dataset. Future "
        "improvements may require additional external features such as ambient temperature, "
        "baseline neighbourhood power consumption, or real-time grid telemetry."
    )
    insert_image(doc, IMG["confusion"], width=Inches(4.2),
                 caption_num=3, caption_title="Confusion Matrix — Random Forest Classifier (Peak Load Risk)")
    add_horizontal_rule(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # 8. FEATURE IMPORTANCE
    # ══════════════════════════════════════════════════════════════════════════
    heading1(doc, "8. Feature Importance Analysis")
    body_para(doc,
        "The Random Forest model provides inherent feature importance scores based on the "
        "mean decrease in Gini impurity across all trees. Figure 4 illustrates the relative "
        "importance of each input feature."
    )
    body_para(doc,
        "The most influential features in descending order of importance are:"
    )
    bullet(doc, "energy_dispensed_kwh — The total energy delivered per session is the strongest predictor of grid stress levels.")
    bullet(doc, "grid_load_mw — Directly reflects the instantaneous grid draw and is naturally correlated with risk level.")
    bullet(doc, "renewable_energy_used_percent — Renewable energy proportion moderates grid stress, making it a meaningful predictor.")
    bullet(doc, "avg_charging_duration_minutes — Longer sessions contribute to sustained grid load.")
    bullet(doc, "hour_of_day — Temporal context captures peak usage periods aligned with human activity schedules.")
    bullet(doc, "vehicles_charged — Higher simultaneous vehicle counts increase aggregate demand.")
    body_para(doc,
        "Notably, no single feature overwhelmingly dominates the prediction. The relatively "
        "uniform distribution of importance scores suggests that peak load risk is a multifactorial "
        "phenomenon driven by the interaction of energy, temporal, and operational variables "
        "simultaneously. This finding reinforces the appropriateness of an ensemble model like "
        "Random Forest over simpler univariate approaches."
    )
    insert_image(doc, IMG["feature_imp"], width=Inches(5.2),
                 caption_num=4, caption_title="Feature Importance Scores — Random Forest Classifier")
    add_horizontal_rule(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # 9. CLASS DISTRIBUTION
    # ══════════════════════════════════════════════════════════════════════════
    heading1(doc, "9. Class Distribution Analysis")
    body_para(doc,
        "Prior to modelling, the distribution of the target variable peak_load_risk was examined "
        "to assess class balance. The class distribution across the full dataset (n = 2,800) is "
        "summarised in Table 6 and visualised in Figure 5."
    )

    table_caption(doc, 6, "Target Variable Class Distribution")
    styled_table(doc,
        headers=["Class", "Frequency", "Percentage"],
        rows=[
            ["Low",    "1,404", "50.1%"],
            ["Medium", "  830", "29.6%"],
            ["High",   "  566", "20.2%"],
            ["Total",  "2,800", "100.0%"],
        ],
        col_widths=[Inches(1.8), Inches(1.5), Inches(1.5)],
    )
    doc.add_paragraph()

    body_para(doc,
        "The dataset exhibits a moderate class imbalance, with the Low-risk class comprising "
        "50.1% of samples compared to only 20.2% for the High-risk class. This imbalance "
        "motivated the application of SMOTE on the training partition to prevent the classifier "
        "from developing a bias towards the majority class. Left unaddressed, such imbalance "
        "would cause the model to neglect the operationally critical High-risk class, resulting "
        "in poor recall for the most dangerous grid events."
    )
    insert_image(doc, IMG["class_dist"], width=Inches(4.8),
                 caption_num=5, caption_title="Class Distribution of Peak Load Risk Target Variable")
    add_horizontal_rule(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # 10. CLUSTERING ANALYSIS
    # ══════════════════════════════════════════════════════════════════════════
    heading1(doc, "10. Clustering Analysis")

    heading2(doc, "10.1 Algorithm and Configuration")
    body_para(doc,
        "K-Means clustering was applied to identify inherent operational groupings within the "
        "charging dataset. The algorithm was configured with k = 3 clusters, informed by the "
        "three-tier risk classification and preliminary elbow curve analysis. Clustering was "
        "performed on the scaled continuous features, primarily energy_dispensed_kwh and "
        "grid_load_mw, to segment stations by their energy and demand profiles."
    )

    heading2(doc, "10.2 Identified Clusters")
    body_para(doc,
        "The K-Means algorithm successfully identified three distinct operational charging "
        "profiles within the dataset:"
    )
    bullet(doc, "Cluster 0 — Low-energy, low-grid-load sessions: Characteristic of slow residential or overnight charging stations with extended durations and minimal grid impact.")
    bullet(doc, "Cluster 1 — Moderate-energy, moderate-grid-load sessions: Representative of standard urban fast chargers serving a mix of commuter and commercial EV users.")
    bullet(doc, "Cluster 2 — High-energy, high-grid-load sessions: Indicative of ultra-fast commercial charging hubs or simultaneous multi-vehicle charging events that place the greatest stress on grid infrastructure.")
    body_para(doc,
        "These clusters provide grid operators with actionable operational archetypes. Rather "
        "than applying a uniform grid limit across all stations, operators can tailor power "
        "allocation and demand-response triggers dynamically based on the operational profile "
        "of each station cluster."
    )

    heading2(doc, "10.3 Cluster Visualisation")
    body_para(doc,
        "Figure 6 presents the K-Means cluster scatter plot, with data points colour-coded "
        "by cluster assignment. The three groups are well-separated in the energy-versus-grid-load "
        "feature space, confirming the validity of k = 3 as an appropriate cluster count for "
        "this dataset."
    )
    # Note: kmeans cluster image is embedded within the dashboard; using dashboard for this
    body_para(doc,
        "(Refer to Panel 2 of the Final Dashboard — Figure 8 — for the K-Means cluster "
        "visualisation across operational zones.)",
        italic=True
    )
    add_horizontal_rule(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # 11. OUTLIER DETECTION
    # ══════════════════════════════════════════════════════════════════════════
    heading1(doc, "11. Outlier Detection")

    heading2(doc, "11.1 Algorithm")
    body_para(doc,
        "Isolation Forest was selected for anomaly detection due to its efficiency with "
        "high-dimensional tabular data and its ability to isolate outliers without making "
        "distributional assumptions. The algorithm works by recursively partitioning the "
        "feature space using random splits; anomalous points require fewer splits to isolate "
        "and therefore receive lower anomaly scores."
    )
    body_para(doc,
        "The model was configured with a contamination rate of 3% (0.03), reflecting the "
        "expectation that a small proportion of sessions represent genuinely anomalous "
        "grid-stress events."
    )

    heading2(doc, "11.2 Results")
    body_para(doc,
        "The Isolation Forest model flagged 84 anomalous charging events from the full "
        "dataset of 2,800 records. These outliers represent charging sessions with unusual "
        "combinations of the following characteristics:"
    )
    bullet(doc, "Disproportionately high grid load (grid_load_mw) relative to the number of vehicles charged.")
    bullet(doc, "Unexpectedly large energy dispensed values relative to session duration.")
    bullet(doc, "Unusual temporal patterns, such as very high loads occurring during typically off-peak hours.")

    heading2(doc, "11.3 Operational Significance")
    body_para(doc,
        "These 84 flagged events represent potential 'black swan' moments for grid stability — "
        "events where the grid load spikes disproportionately to the apparent charging demand. "
        "Such events may arise from:"
    )
    bullet(doc, "Hardware malfunctions at specific charging stations causing erroneous power draw.")
    bullet(doc, "Simultaneous large-scale charging events coinciding with baseline grid stress.")
    bullet(doc, "Data recording anomalies or sensor faults at station level.")
    body_para(doc,
        "Identifying and flagging these events allows grid engineers to investigate specific "
        "stations, prepare emergency battery buffer reserves, and implement real-time circuit "
        "breaker protocols for similar future events."
    )
    body_para(doc,
        "(Refer to Panel 3 of the Final Dashboard — Figure 8 — for the Isolation Forest "
        "anomaly visualisation.)",
        italic=True
    )
    add_horizontal_rule(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # 12. DASHBOARD VISUALISATION
    # ══════════════════════════════════════════════════════════════════════════
    heading1(doc, "12. Dashboard Visualisation")
    body_para(doc,
        "A comprehensive three-panel dashboard (Figure 8) was developed to consolidate the "
        "key analytical findings into a single visual summary suitable for operational "
        "decision-making by grid managers and EV infrastructure planners."
    )

    heading2(doc, "Panel 1: Temporal Risk — Hourly Grid Load")
    body_para(doc,
        "The first panel presents a temporal visualisation of grid load aggregated by hour "
        "of the day. The plot reveals significant intraday volatility, with pronounced spikes "
        "during morning commute hours (07:00–09:00) and evening charging peaks (18:00–21:00). "
        "Deep drops in grid load are observed during overnight hours, consistent with reduced "
        "EV usage. This temporal analysis supports the implementation of time-of-use pricing "
        "strategies to flatten the demand curve and reduce peak grid stress."
    )

    heading2(doc, "Panel 2: Operational Clusters — K-Means")
    body_para(doc,
        "The second panel visualises the K-Means cluster assignments in the two-dimensional "
        "energy-versus-grid-load feature space. Three clearly delineated clusters are evident, "
        "demonstrating the algorithm's ability to partition the operational space into "
        "meaningfully distinct charging profiles. These clusters can be directly mapped to "
        "station-level management policies."
    )

    heading2(doc, "Panel 3: Anomaly Detection — Isolation Forest")
    body_para(doc,
        "The third panel overlays the Isolation Forest anomaly flags on the scatter plot "
        "of charging sessions. The 84 flagged anomalous events (visualised in red) are "
        "clearly distinguishable from the 2,716 normal sessions (in blue), confirming that "
        "the algorithm has successfully isolated a small, operationally significant set of "
        "unusual charging events."
    )
    insert_image(doc, IMG["dashboard"], width=Inches(6.2),
                 caption_num=6, caption_title="Final Three-Panel Analysis Dashboard: (1) Hourly Grid Load, (2) K-Means Clusters, (3) Isolation Forest Anomalies")
    add_horizontal_rule(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # 13. RESULTS AND DISCUSSION
    # ══════════════════════════════════════════════════════════════════════════
    heading1(doc, "13. Results and Discussion")

    heading2(doc, "13.1 Correlation Analysis Findings")
    body_para(doc,
        "The Pearson correlation analysis revealed predominantly weak linear relationships "
        "among all numerical features. This result, while at first appearing to limit predictive "
        "potential, is an informative finding in itself: it indicates that peak load risk is not "
        "driven by any single dominant linear relationship, but rather emerges from complex, "
        "possibly non-linear interactions among multiple operational variables simultaneously. "
        "This insight motivated the adoption of an ensemble Random Forest approach rather than "
        "a simple linear classifier."
    )

    heading2(doc, "13.2 City Zone Similarity and Dissimilarity")
    body_para(doc,
        "The Euclidean distance analysis revealed that South and West zones are operationally "
        "the most similar (distance ≈ 0.85), while East and South are the most dissimilar "
        "(distance ≈ 3.64). This finding has direct policy implications: grid management "
        "protocols that are effective in the South zone may be directly applicable to the West "
        "zone with minimal adaptation, whereas the East zone requires a distinctly tailored "
        "operational strategy. Zone-specific demand forecasting models would better capture "
        "this geographic heterogeneity than a single global model."
    )

    heading2(doc, "13.3 Classification Performance")
    body_para(doc,
        "The Random Forest Classifier achieved an overall accuracy of 45.18% — a result that "
        "is modest but contextually meaningful. The three-class peak load risk prediction "
        "problem is inherently challenging given the overlapping feature distributions between "
        "risk classes, the weak linear separability of features, and the limitations of the "
        "available feature set. The SMOTE technique partially mitigated the class imbalance "
        "effect, improving recall for minority classes at the cost of some precision. Low-risk "
        "events were classified most reliably (F1 = 0.58), while High-risk events (F1 = 0.25) "
        "remain the most operationally critical gap in the current model."
    )

    heading2(doc, "13.4 Feature Importance")
    body_para(doc,
        "The feature importance analysis confirmed that energy_dispensed_kwh and grid_load_mw "
        "are the most predictive features, which is intuitively consistent with the physical "
        "nature of grid stress. The relatively uniform distribution of importance scores across "
        "features reinforces the multifactorial nature of peak load risk and suggests that "
        "future models should pursue richer feature sets rather than more sophisticated algorithms."
    )

    heading2(doc, "13.5 Clustering Insights")
    body_para(doc,
        "K-Means clustering successfully partitioned the dataset into three well-separated "
        "operational clusters, demonstrating that meaningful structure exists within the data "
        "even when supervised classification performance is limited. These clusters offer an "
        "alternative, unsupervised lens through which grid operators can categorise stations "
        "and allocate resources without relying on explicit risk labels."
    )

    heading2(doc, "13.6 Anomaly Detection")
    body_para(doc,
        "The Isolation Forest model identified 84 anomalous events, representing approximately "
        "3% of all charging sessions. These outliers exhibit unusual combinations of high grid "
        "load and charging demand that may represent hardware faults, data quality issues, or "
        "genuine grid stress events. Proactively flagging and investigating these events is "
        "essential for maintaining grid reliability and preventing cascading infrastructure failures."
    )
    add_horizontal_rule(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # 14. CONCLUSION
    # ══════════════════════════════════════════════════════════════════════════
    heading1(doc, "14. Conclusion")
    body_para(doc,
        "This case study successfully demonstrated the application of a comprehensive data "
        "mining pipeline to an EV charging station dataset. Beginning with thorough data "
        "preprocessing — including temporal feature engineering, one-hot encoding, target "
        "variable mapping, and feature scaling — the study established a clean, analysis-ready "
        "dataset of (2,800, 15) dimensions."
    )
    body_para(doc,
        "The similarity analysis revealed that numerical features exhibit predominantly weak "
        "linear correlations, while the Euclidean distance analysis identified meaningful "
        "operational dissimilarities across city zones — particularly between the East and "
        "South zones. These findings inform zone-specific grid management strategies."
    )
    body_para(doc,
        "The Random Forest Classifier achieved an overall accuracy of 45.18% for the three-class "
        "peak load risk prediction task. While moderate, this result is consistent with the weak "
        "feature correlations and the synthetic nature of the dataset. Feature importance analysis "
        "confirmed that energy_dispensed_kwh, grid_load_mw, and renewable_energy_used_percent are "
        "the most predictive variables, and that peak load risk is a genuinely multifactorial "
        "phenomenon."
    )
    body_para(doc,
        "K-Means clustering (k = 3) successfully identified three distinct operational charging "
        "profiles, providing a practical segmentation framework for targeted station-level "
        "resource management. Isolation Forest anomaly detection flagged 84 anomalous events "
        "representing potential high-risk grid stress scenarios requiring operational investigation."
    )
    body_para(doc,
        "The integrated analytical framework presented in this study offers a foundation for "
        "intelligent EV infrastructure planning. Future work should focus on incorporating "
        "external macroeconomic, weather, and real-time grid telemetry data to improve "
        "classification accuracy and support the development of dynamic, data-driven grid "
        "management systems."
    )
    add_horizontal_rule(doc)

    # ══════════════════════════════════════════════════════════════════════════
    # REFERENCES
    # ══════════════════════════════════════════════════════════════════════════
    heading1(doc, "References")

    references = [
        "[1] Kaggle. (2024). EV Charging Station Usage and Grid Load Analysis Dataset. "
            "Retrieved from https://www.kaggle.com. "
            "Licensed under Kaggle Open Data.",

        "[2] Scikit-Learn Developers. (2024). scikit-learn: Machine Learning in Python. "
            "sklearn.ensemble.RandomForestClassifier, sklearn.cluster.KMeans, "
            "sklearn.ensemble.IsolationForest, sklearn.preprocessing.StandardScaler. "
            "Available at: https://scikit-learn.org/stable/",

        "[3] McKinney, W. (2010). Data Structures for Statistical Computing in Python. "
            "Proceedings of the 9th Python in Science Conference. "
            "Pandas Documentation: https://pandas.pydata.org/docs/",

        "[4] Breiman, L. (2001). Random Forests. Machine Learning, 45(1), 5–32. "
            "https://doi.org/10.1023/A:1010933404324",

        "[5] MacQueen, J. (1967). Some Methods for Classification and Analysis of Multivariate "
            "Observations. Proceedings of the 5th Berkeley Symposium on Mathematical Statistics "
            "and Probability, 1, 281–297.",

        "[6] Liu, F. T., Ting, K. M., & Zhou, Z.-H. (2008). Isolation Forest. "
            "2008 Eighth IEEE International Conference on Data Mining, 413–422. "
            "https://doi.org/10.1109/ICDM.2008.17",

        "[7] Chawla, N. V., Bowyer, K. W., Hall, L. O., & Kegelmeyer, W. P. (2002). "
            "SMOTE: Synthetic Minority Over-sampling Technique. "
            "Journal of Artificial Intelligence Research, 16, 321–357. "
            "https://doi.org/10.1613/jair.953",

        "[8] Hunter, J. D. (2007). Matplotlib: A 2D Graphics Environment. "
            "Computing in Science & Engineering, 9(3), 90–95. "
            "https://doi.org/10.1109/MCSE.2007.55",

        "[9] Waskom, M. L. (2021). Seaborn: Statistical Data Visualization. "
            "Journal of Open Source Software, 6(60), 3021. "
            "https://doi.org/10.21105/joss.03021",
    ]

    for ref in references:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        p.paragraph_format.left_indent  = Cm(0.8)
        p.paragraph_format.first_line_indent = Cm(-0.8)
        run = p.add_run(ref)
        run.font.size = Pt(10.5)
        run.font.color.rgb = BODY_COLOR

    # ── save ──────────────────────────────────────────────────────────────────
    doc.save(OUTPUT)
    print(f"\n[OK] Report successfully generated:\n   {OUTPUT}\n")


if __name__ == "__main__":
    build_report()
